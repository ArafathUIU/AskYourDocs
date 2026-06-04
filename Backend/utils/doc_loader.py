from pathlib import Path
import re


def extract_text_from_docx(docx_path: Path) -> tuple[str, dict]:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for DOCX support. Install: pip install python-docx")

    doc = Document(str(docx_path))
    full_text = []
    page_count = len(doc.paragraphs) // 40 or 1  # rough estimate

    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())

    text = "\n\n".join(full_text)
    text = f"[Page 1]\n{text}"

    name = docx_path.stem
    if doc.paragraphs:
        first = doc.paragraphs[0].text.strip()
        if first:
            name = first[:100]

    return text, {
        "title": name,
        "author": doc.core_properties.author or "",
        "page_count": page_count,
    }


def extract_text_from_txt(txt_path: Path) -> tuple[str, dict]:
    with open(txt_path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()

    # Estimate pages (roughly 3000 chars per page)
    page_count = max(1, len(text) // 3000)
    text = f"[Page 1]\n{text.strip()}"

    return text, {
        "title": txt_path.stem,
        "author": "",
        "page_count": page_count,
    }


def extract_text_from_md(md_path: Path) -> tuple[str, dict]:
    with open(md_path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()

    page_count = max(1, len(text) // 3000)
    text = f"[Page 1]\n{text.strip()}"

    title = md_path.stem
    h1_match = re.search(r'^#\s+(.+)$', text, re.MULTILINE)
    if h1_match:
        title = h1_match.group(1)

    return text, {
        "title": title,
        "author": "",
        "page_count": page_count,
    }


SUPPORTED_EXTENSIONS = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".txt": "txt",
    ".md": "md",
    ".markdown": "md",
}


def extract_text_from_file(file_path: Path) -> tuple[str, dict]:
    ext = file_path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {ext}")

    if ext == ".pdf":
        from utils.pdf_loader import extract_text_from_pdf, get_pdf_metadata
        text = extract_text_from_pdf(file_path)
        meta = get_pdf_metadata(file_path)
        return text, meta

    if ext == ".docx":
        return extract_text_from_docx(file_path)

    if ext in (".txt",):
        return extract_text_from_txt(file_path)

    if ext in (".md", ".markdown"):
        return extract_text_from_md(file_path)


def get_supported_extensions() -> list[str]:
    return list(SUPPORTED_EXTENSIONS.keys())
